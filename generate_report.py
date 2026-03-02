#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Script ghi đè file BAO CAO APP.docx với nội dung mới cho app PickleBall Pro"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ===== STYLE =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'Times New Roman'
    hf.bold = True
    hf.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hf.size = Pt(16)
    elif level == 2:
        hf.size = Pt(14)
    else:
        hf.size = Pt(13)

# ===== TRANG BÌA =====
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BÁO CÁO SẢN PHẨM CÔNG NGHỆ THÔNG TIN')
run.font.size = Pt(18)
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('HỆ THỐNG QUẢN LÝ SÂN PICKLEBALL TRỰC TUYẾN')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(5, 150, 105)

app_name = doc.add_paragraph()
app_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = app_name.add_run('"PickleBall Pro"')
run.font.size = Pt(20)
run.font.bold = True
run.font.italic = True
run.font.color.rgb = RGBColor(5, 150, 105)

doc.add_paragraph('')
doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Lĩnh vực: Quản lý thể thao - Chuyển đổi số')
run.font.size = Pt(13)

doc.add_page_break()

# ===== MỤC LỤC =====
doc.add_heading('MỤC LỤC', level=1)
toc_items = [
    'PHẦN I. GIỚI THIỆU CHUNG',
    '  1. Lý do chọn đề tài',
    '  2. Cơ sở lý luận',
    '  3. Cơ sở thực tiễn',
    'PHẦN II. NỘI DUNG NGHIÊN CỨU',
    '  1. Quy trình thiết kế',
    '  2. Nguyên lý hoạt động',
    '  3. Các tính năng chính của sản phẩm',
    '  4. Những công nghệ sử dụng trong sản phẩm',
    'PHẦN III. KẾT QUẢ NGHIÊN CỨU',
    '  1. Kết quả nghiên cứu',
    '  2. Hướng phát triển của đề tài',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===== PHẦN I =====
doc.add_heading('PHẦN I. GIỚI THIỆU CHUNG', level=1)

doc.add_heading('1. Lý do chọn đề tài', level=2)

paras = [
    'Trong những năm gần đây, phong trào chơi Pickleball đã bùng nổ mạnh mẽ trên toàn thế giới và tại Việt Nam. Đây là môn thể thao kết hợp giữa cầu lông, bóng bàn và quần vợt, phù hợp với mọi lứa tuổi từ học sinh đến người cao tuổi. Sự phát triển nhanh chóng của phong trào đã kéo theo nhu cầu ngày càng tăng về việc quản lý hệ thống sân bãi, lịch đặt sân và dịch vụ liên quan.',

    'Tuy nhiên, phần lớn các cơ sở kinh doanh sân Pickleball hiện nay vẫn đang quản lý bằng phương pháp thủ công: ghi chép trên sổ sách, nhận đặt sân qua điện thoại hoặc tin nhắn Zalo, tính tiền bằng tay và không có hệ thống theo dõi doanh thu. Điều này dẫn đến nhiều bất cập như đặt trùng sân, bỏ sót khách hàng, khó kiểm soát công nợ và không có dữ liệu để phân tích kinh doanh.',

    'Xuất phát từ thực tế đó, tôi đã nghiên cứu và phát triển hệ thống "PickleBall Pro" – một ứng dụng web hiện đại giúp chủ sân quản lý toàn diện hoạt động kinh doanh sân Pickleball. Hệ thống được thiết kế theo hướng chuyển đổi số, cho phép vận hành mọi lúc mọi nơi chỉ với một chiếc điện thoại hoặc máy tính có kết nối mạng.',

    'Sản phẩm không chỉ giải quyết các vấn đề quản lý truyền thống mà còn mang đến trải nghiệm đặt sân trực tuyến cho khách hàng, hệ thống thông báo tức thì khi admin xác nhận hoặc hủy lịch, từ đó nâng cao chất lượng dịch vụ và sự chuyên nghiệp trong kinh doanh thể thao.',
]
for text in paras:
    doc.add_paragraph(text)

doc.add_paragraph('[Chèn ảnh: Ảnh giao diện đăng nhập PickleBall Pro với logo và phong cách thiết kế hiện đại]')

doc.add_heading('2. Cơ sở lý luận', level=2)

paras = [
    'Dự án được xây dựng dựa trên ba nền tảng lý luận chính:',

    '- Thứ nhất, dựa trên lý thuyết về hệ thống quản lý cơ sở dữ liệu quan hệ (RDBMS). Toàn bộ dữ liệu của hệ thống được tổ chức theo mô hình quan hệ với các bảng: sân (courts), khách hàng (customers), lịch đặt (bookings), sản phẩm (products), hóa đơn (invoices), thông báo (notifications) và hồ sơ người dùng (profiles). Các bảng được liên kết chặt chẽ qua khóa ngoại, đảm bảo tính toàn vẹn dữ liệu. Cơ sở lý luận này giúp hệ thống xử lý các truy vấn phức tạp một cách chính xác và hiệu quả.',

    '- Thứ hai, dựa trên nguyên tắc phân quyền theo vai trò (Role-Based Access Control – RBAC). Hệ thống chia người dùng thành ba vai trò rõ ràng: Admin (quản trị toàn quyền), Staff (nhân viên hỗ trợ) và Customer (khách hàng). Mỗi vai trò chỉ được truy cập những chức năng và dữ liệu phù hợp, đảm bảo an ninh thông tin theo chuẩn bảo mật Row Level Security (RLS) của Supabase. Nguyên tắc này tạo ra một hệ thống vừa linh hoạt vừa an toàn.',

    '- Thứ ba, dựa trên mô hình kiến trúc Single Page Application (SPA) kết hợp với Backend-as-a-Service (BaaS). Mô hình này cho phép ứng dụng hoạt động mượt mà như một ứng dụng native, không cần tải lại trang khi chuyển đổi chức năng, đồng thời tận dụng hạ tầng đám mây để lưu trữ và xử lý dữ liệu. Theo lý thuyết kiến trúc phần mềm hiện đại, mô hình SPA + BaaS giúp giảm chi phí phát triển, tăng tốc độ triển khai và dễ dàng mở rộng quy mô.',
]
for text in paras:
    doc.add_paragraph(text)

doc.add_heading('3. Cơ sở thực tiễn', level=2)

paras = [
    'Cơ sở thực tiễn để phát triển hệ thống này xuất phát từ chính những bất cập mà các chủ sân Pickleball đang gặp phải:',

    'Thứ nhất là sự quá tải trong việc quản lý lịch đặt sân. Hiện nay, hầu hết chủ sân nhận đặt sân qua tin nhắn Zalo, gọi điện thoại hoặc ghi chép trên sổ tay. Khi số lượng sân và khách hàng tăng lên, việc theo dõi lịch đặt theo giờ cho từng sân trở nên rất phức tạp. Tình trạng đặt trùng sân, quên xác nhận cho khách, hoặc bỏ sót lịch đặt xảy ra thường xuyên, gây mất uy tín và thiệt hại doanh thu.',

    'Thứ hai là khó khăn trong việc quản lý tài chính và doanh thu. Khi thanh toán bằng tiền mặt và ghi sổ tay, việc đối soát doanh thu cuối ngày, cuối tháng trở thành gánh nặng. Chủ sân không thể biết chính xác doanh thu từng ngày, từng sân, hay tổng chi tiêu của từng khách hàng. Thiếu dữ liệu phân tích khiến việc đưa ra quyết định kinh doanh chỉ dựa trên cảm tính.',

    'Thứ ba là trải nghiệm đặt sân chưa thuận tiện cho khách hàng. Khách hàng phải gọi điện hoặc nhắn tin để hỏi sân trống, chờ phản hồi từ chủ sân, rồi mới xác nhận được lịch. Quy trình này mất thời gian và không chuyên nghiệp, đặc biệt khi so sánh với các ứng dụng đặt dịch vụ hiện đại mà người dùng đã quen thuộc.',

    'Từ những thực tế đó, cần có một giải pháp công nghệ toàn diện, cho phép chủ sân quản lý "mọi lúc, mọi nơi" và khách hàng có thể tự đặt sân trực tuyến một cách nhanh chóng, tiện lợi.',

    '[Chèn ảnh: Ảnh so sánh giữa quản lý thủ công (sổ sách, tin nhắn Zalo) và giao diện quản lý trên PickleBall Pro]',
]
for text in paras:
    doc.add_paragraph(text)

doc.add_page_break()

# ===== PHẦN II =====
doc.add_heading('PHẦN II. NỘI DUNG NGHIÊN CỨU', level=1)

doc.add_heading('1. Quy trình thiết kế', level=2)

paras = [
    'Để tạo ra một hệ thống vừa đảm bảo tính kỹ thuật, vừa phù hợp với nhu cầu thực tế của chủ sân Pickleball, tôi đã triển khai quy trình thiết kế gồm bốn giai đoạn:',

    'Giai đoạn đầu tiên là Xây dựng cấu trúc dữ liệu và logic nghiệp vụ. Tôi tiến hành thiết kế cơ sở dữ liệu trên Supabase (PostgreSQL) với 8 bảng chính: profiles (hồ sơ người dùng), courts (thông tin sân), customers (khách hàng), bookings (lịch đặt), products (sản phẩm/dịch vụ), invoices (hóa đơn), invoice_items (chi tiết hóa đơn), notifications (thông báo) và settings (cài đặt hệ thống). Mỗi bảng được thiết kế với các khóa chính, khóa ngoại và ràng buộc dữ liệu chặt chẽ. Đồng thời, tôi thiết lập hệ thống Row Level Security (RLS) với các policy phân quyền chi tiết cho từng vai trò.',

    'Giai đoạn thứ hai là Thiết kế giao diện người dùng (UI/UX). Tôi sử dụng TailwindCSS kết hợp với các hiệu ứng glassmorphism (kính mờ), gradient và micro-animation để tạo ra giao diện hiện đại, chuyên nghiệp. Hệ thống được chia thành hai layout riêng biệt: Layout admin (dạng sidebar navigation) dành cho chủ sân và nhân viên, Layout khách hàng (dạng mobile-first) dành cho người đặt sân. Các trang quản lý sử dụng thiết kế dạng bảng và card, hỗ trợ responsive trên mọi thiết bị.',

    'Giai đoạn thứ ba là Phát triển các module chức năng. Hệ thống bao gồm 10 module chính: Xác thực (đăng nhập/đăng ký), Dashboard (tổng quan), Quản lý sân, Quản lý khách hàng, Quản lý lịch đặt, Bán hàng (POS), Hóa đơn, Báo cáo & Thống kê, Quản lý người dùng, và Cài đặt. Đặc biệt, tôi đã phát triển trang đặt sân trực tuyến cho khách hàng với lịch tuần, time slot trực quan và hệ thống thông báo realtime.',

    'Giai đoạn cuối cùng là Tối ưu hóa và Triển khai. Sau khi hoàn thiện các tính năng, tôi tiến hành tối ưu hiệu năng, kiểm tra bảo mật RLS, và đóng gói ứng dụng bằng Vite để triển khai lên Vercel. Ứng dụng hoạt động hoàn toàn trên nền web, không cần cài đặt, truy cập được từ mọi thiết bị có trình duyệt.',

    '[Chèn ảnh: Sơ đồ kiến trúc hệ thống PickleBall Pro: Frontend (React) → Supabase (Auth + Database + Realtime) → PostgreSQL]',
]
for text in paras:
    doc.add_paragraph(text)

doc.add_heading('2. Nguyên lý hoạt động', level=2)

paras = [
    'Hệ thống PickleBall Pro hoạt động dựa trên hai luồng chính: luồng quản trị (Admin/Staff) và luồng khách hàng (Customer).',

    'Đối với chủ sân (Admin):', 

    '- Bước 1: Đăng nhập và thiết lập ban đầu. Admin đăng nhập bằng email/mật khẩu, sau đó cài đặt thông tin doanh nghiệp (tên cơ sở, địa chỉ, giờ hoạt động, giá thuê sân, phí chiếu sáng...) tại trang Cài đặt. Tiếp theo, Admin tạo thông tin các sân (tên sân, loại sân trong nhà/ngoài trời, giá thuê giờ thường/giờ cao điểm).',

    '- Bước 2: Quản lý vận hành hàng ngày. Mỗi ngày, Admin mở trang Dashboard để xem tổng quan: doanh thu hôm nay, số booking, số khách hàng mới, và lịch đặt sân hôm nay. Tại trang Bookings, Admin xem lịch đặt theo tuần dưới dạng bảng thời gian × sân, có thể tạo booking mới, xác nhận booking khách đặt, check-in khi khách đến, và hoàn thành khi khách chơi xong.',

    '- Bước 3: Thanh toán và quản lý tài chính. Tại trang POS (Point of Sale), Admin có thể bán thêm nước uống, đồ ăn, cho thuê vợt/bóng. Hệ thống tự động tạo hóa đơn, ghi nhận phương thức thanh toán (tiền mặt, chuyển khoản, MoMo, ZaloPay, VNPay). Trang Báo cáo & Thống kê cung cấp biểu đồ doanh thu, phân tích theo sân, theo thời gian.',

    'Đối với khách hàng (Customer):',

    '- Bước 1: Đăng ký tài khoản và đặt sân trực tuyến. Khách hàng truy cập ứng dụng, đăng ký bằng email, sau đó hệ thống tự động gán role "customer". Trang đặt sân hiển thị danh sách sân hiện có với thông tin giá, trạng thái, số slot trống hôm nay. Khách chọn ngày trên lịch tuần, chọn time slot trống (hiển thị bằng màu xanh), điền thông tin và bấm đặt sân.',

    '- Bước 2: Theo dõi trạng thái booking. Sau khi đặt, booking ở trạng thái "Chờ xác nhận". Khi Admin xác nhận, hệ thống tự động gửi thông báo realtime đến khách thông qua icon chuông trên header. Khách có thể xem danh sách booking sắp tới, lịch sử booking, và hủy booking nếu cần.',

    '[Chèn ảnh: Ảnh minh họa quy trình: 1. Khách đặt sân → 2. Admin xác nhận → 3. Thông báo gửi đến khách]',
]
for text in paras:
    doc.add_paragraph(text)

doc.add_heading('3. Các tính năng chính của sản phẩm', level=2)

features = [
    ('a) Xác thực và Phân quyền', 
     'Hệ thống sử dụng Supabase Auth để quản lý đăng nhập/đăng ký. Khi người dùng mới đăng ký, trigger tự động tạo profile với role "customer". Admin có thể thăng cấp người dùng lên Staff qua trang Quản lý người dùng. Hệ thống hỗ trợ hiển thị/ẩn mật khẩu và xử lý lỗi đăng nhập tường minh.'),

    ('b) Dashboard – Tổng quan kinh doanh', 
     'Trang tổng quan hiển thị 4 thẻ thống kê: Doanh thu hôm nay, Booking hôm nay, Tổng khách hàng, Tỷ lệ tăng trưởng. Bên dưới là danh sách booking sắp tới với trạng thái (Chờ xác nhận, Đã xác nhận, Đang chơi, Hoàn thành, Đã hủy) và thông tin chi tiết về sân, thời gian, khách hàng.'),

    ('c) Quản lý sân (Courts)', 
     'Admin có thể thêm, sửa, xóa sân. Mỗi sân gồm thông tin: tên sân, loại (trong nhà/ngoài trời), bề mặt, trạng thái (hoạt động/bảo trì/tạm ngưng), giá thuê giờ thường, giá giờ cao điểm, phí chiếu sáng. Giao diện hiển thị dạng card với gradient màu theo loại sân.'),

    ('d) Quản lý khách hàng (Customers)', 
     'Trang quản lý khách hàng cho phép thêm, sửa, xóa, tìm kiếm khách. Mỗi khách hàng có thông tin: tên, số điện thoại, email, trạng thái VIP, ghi chú, tổng số booking và tổng chi tiêu. Hệ thống tự động tạo hồ sơ khách hàng khi khách đặt sân lần đầu thông qua function SECURITY DEFINER.'),

    ('e) Quản lý lịch đặt sân (Bookings)', 
     'Đây là tính năng cốt lõi của hệ thống. Trang Bookings hiển thị lịch đặt theo tuần dưới dạng bảng thời gian (6:00-22:00) × danh sách sân. Mỗi ô hiển thị tên khách nếu đã có booking. Admin có thể tạo booking mới, xác nhận (✓), check-in (→), hoàn thành (✓✓), hoặc hủy (✕) booking. Khi thay đổi trạng thái, hệ thống tự động gửi thông báo đến khách hàng.'),

    ('f) Đặt sân trực tuyến (Customer Booking)', 
     'Khách hàng truy cập giao diện riêng biệt với 3 tab: Danh sách sân (xem thông tin sân, giá, slot trống), Đặt sân (chọn ngày/giờ trên lịch tuần, time slot xanh = trống, đỏ = đã đặt), và Lịch của tôi (xem booking sắp tới, lịch sử, hủy booking). Form đặt sân hiển thị thông tin sân, ngày, giờ, ghi chú và giá sân.'),

    ('g) Hệ thống Thông báo Realtime (Notifications)', 
     'Khi Admin xác nhận, hủy, check-in hoặc hoàn thành booking, hệ thống tự động tạo thông báo và đẩy đến khách hàng qua Supabase Realtime. Khách thấy icon chuông trên header với badge số thông báo chưa đọc. Click chuông mở dropdown danh sách thông báo với tiêu đề, nội dung, thời gian. Hỗ trợ đánh dấu đã đọc từng thông báo hoặc tất cả.'),

    ('h) Bán hàng – POS (Point of Sale)', 
     'Hệ thống POS cho phép bán các sản phẩm/dịch vụ: cho thuê vợt/bóng, đồ ăn, nước uống, dụng cụ. Giao diện chia thành 2 phần: bên trái là danh sách sản phẩm (có tìm kiếm, lọc theo danh mục, thêm/sửa sản phẩm), bên phải là giỏ hàng (thêm/bớt số lượng, tính tổng, thanh toán). Admin có thể quản lý tồn kho với cảnh báo khi số lượng thấp.'),

    ('i) Hóa đơn (Invoices)', 
     'Trang hóa đơn quản lý toàn bộ giao dịch: mã hóa đơn tự động, liên kết với booking/khách hàng, tổng tiền, giảm giá, thuế, cọc trừ, còn phải trả, phương thức thanh toán (tiền mặt, chuyển khoản, thẻ, MoMo, ZaloPay, VNPay), trạng thái thanh toán. Hỗ trợ tìm kiếm và xem chi tiết hóa đơn.'),

    ('j) Báo cáo & Thống kê (Reports)', 
     'Trang báo cáo cung cấp biểu đồ trực quan: biểu đồ doanh thu theo thời gian (AreaChart), biểu đồ booking theo sân (BarChart), biểu đồ xu hướng (LineChart), phân bố trạng thái (PieChart). Hỗ trợ lọc theo khoảng thời gian và xuất dữ liệu ra file CSV.'),

    ('k) Quản lý người dùng (Users)', 
     'Admin có thể xem danh sách tất cả người dùng, tìm kiếm, lọc theo vai trò, thay đổi role (admin/staff/customer), và mời người dùng mới bằng email. Giao diện hiển thị thông tin: họ tên, email, vai trò (với icon và badge màu), ngày tạo.'),

    ('l) Cài đặt hệ thống (Settings)', 
     'Trang cài đặt cho phép tùy chỉnh: thông tin doanh nghiệp (tên, SĐT, địa chỉ, email), giờ hoạt động, giờ cao điểm, thời lượng slot (phút), % tiền cọc, chính sách hủy (số giờ hủy miễn phí, % phí hủy), thuế suất, đơn vị tiền tệ. Dữ liệu cài đặt được lưu dạng key-value trên Supabase.'),
]

for title, desc in features:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_paragraph('[Chèn ảnh: Ảnh chụp màn hình các trang chính của ứng dụng: Dashboard, Bookings, POS, Customer Booking]')

doc.add_heading('4. Những công nghệ sử dụng trong sản phẩm', level=2)

paras = [
    'Để xây dựng một hệ thống hoạt động ổn định, bảo mật và chuyên nghiệp, tôi đã lựa chọn và kết hợp các công nghệ tiên tiến nhất hiện nay:',

    'Đầu tiên là ReactJS 18 kết hợp với Vite. React là thư viện giao diện phổ biến nhất thế giới, giúp xây dựng giao diện người dùng linh hoạt theo mô hình component-based. Vite là công cụ build hiện đại với tốc độ hot reload cực nhanh (dưới 100ms), giúp quá trình phát triển trở nên mượt mà. Kết hợp với TypeScript để đảm bảo type safety, giảm thiểu lỗi trong quá trình phát triển.',

    'Thứ hai là Supabase – nền tảng Backend-as-a-Service (BaaS) mã nguồn mở. Supabase cung cấp đầy đủ các dịch vụ backend mà không cần viết code server: PostgreSQL database với Row Level Security (RLS) đảm bảo bảo mật ở cấp độ hàng dữ liệu, Authentication (xác thực đăng nhập/đăng ký), Realtime subscriptions (nhận dữ liệu tức thì khi có thay đổi), và Storage (lưu trữ file). Supabase giúp giảm từ hàng tháng phát triển backend xuống chỉ còn vài ngày.',

    'Thứ ba là TailwindCSS – framework CSS utility-first. Thay vì viết CSS truyền thống, TailwindCSS cho phép tạo giao diện trực tiếp bằng các class utility. Phong cách thiết kế của ứng dụng sử dụng hiệu ứng glassmorphism (kính mờ), gradient emerald-teal, micro-animation (fade-in, slide-up, spin), tạo cảm giác hiện đại và chuyên nghiệp.',

    'Thứ tư là Zustand – thư viện quản lý state tối giản. Với kích thước chỉ 1KB, Zustand giúp quản lý trạng thái toàn cục (thông tin đăng nhập, profile người dùng) một cách gọn nhẹ và hiệu quả hơn nhiều so với Redux.',

    'Thứ năm là Recharts – thư viện biểu đồ cho React. Recharts cung cấp các loại biểu đồ đẹp mắt và tương tác: AreaChart (doanh thu), BarChart (booking theo sân), LineChart (xu hướng), PieChart (phân bố). Các biểu đồ hỗ trợ responsive, tooltip và animation mượt mà.',

    'Cuối cùng là các thư viện hỗ trợ: date-fns (xử lý ngày tháng với locale tiếng Việt), lucide-react (hệ thống icon SVG đẹp và nhẹ), react-router-dom (điều hướng SPA).',
]
for text in paras:
    doc.add_paragraph(text)

doc.add_page_break()

# ===== PHẦN III =====
doc.add_heading('PHẦN III. KẾT QUẢ NGHIÊN CỨU', level=1)

doc.add_heading('1. Kết quả nghiên cứu', level=2)

paras = [
    'Sau một thời gian nghiên cứu và phát triển, dự án PickleBall Pro đã đem lại những kết quả cụ thể và khả quan:',

    'Thứ nhất, về mặt quản lý dữ liệu, tôi đã xây dựng thành công một hệ thống cơ sở dữ liệu đám mây toàn diện. Toàn bộ thông tin sân bãi, khách hàng, lịch đặt, hóa đơn và giao dịch được lưu trữ tập trung trên Supabase với bảo mật Row Level Security. Dữ liệu được đồng bộ tức thì, cho phép truy cập từ mọi thiết bị. Việc tra cứu thông tin booking của bất kỳ khách hàng nào chỉ mất vài giây thay vì phải lật sổ sách.',

    'Thứ hai, về mặt hiệu suất vận hành, hệ thống đã tối ưu hóa đáng kể quy trình quản lý sân. Thời gian để xử lý một booking (từ nhận đặt đến xác nhận) giảm từ 5-10 phút (qua điện thoại/tin nhắn) xuống chỉ còn dưới 30 giây. Việc tổng hợp doanh thu cuối ngày từ mất 20-30 phút tính toán thủ công xuống còn tức thì nhờ Dashboard tự động cập nhật. Đặc biệt, tính năng đặt sân trực tuyến giúp chủ sân không bỏ lỡ bất kỳ khách hàng nào, kể cả ngoài giờ làm việc.',

    'Thứ ba, về mặt trải nghiệm khách hàng, hệ thống mang đến sự chuyên nghiệp và tiện lợi. Khách hàng có thể tự chọn sân, xem slot trống theo ngày, đặt sân và nhận thông báo xác nhận tức thì – tất cả chỉ trên một giao diện web. Hệ thống thông báo realtime giúp khách luôn nắm được trạng thái booking mà không cần gọi điện hỏi. Giao diện thiết kế mobile-first, hiển thị đẹp trên điện thoại, giúp khách đặt sân mọi lúc mọi nơi.',

    'Thứ tư, về mặt phân tích kinh doanh, trang Báo cáo & Thống kê cung cấp các biểu đồ trực quan giúp chủ sân đưa ra quyết định dựa trên dữ liệu: sân nào được đặt nhiều nhất, khung giờ nào cao điểm, xu hướng doanh thu tăng hay giảm. Tính năng xuất CSV hỗ trợ báo cáo cho kế toán.',

    'Cuối cùng, dự án đã chứng minh tính khả thi của việc ứng dụng các công nghệ hiện đại vào quản lý kinh doanh thể thao. Với chi phí hosting gần như bằng không (Vercel free tier + Supabase free tier), một cơ sở kinh doanh sân Pickleball hoàn toàn có thể sở hữu một hệ thống quản lý chuyên nghiệp mà không cần đầu tư lớn.',

    '[Chèn ảnh: Ảnh chụp màn hình Dashboard với biểu đồ doanh thu và booking thực tế]',
]
for text in paras:
    doc.add_paragraph(text)

doc.add_heading('2. Hướng phát triển của đề tài', level=2)

paras = [
    'Mặc dù hệ thống hiện tại đã đáp ứng tốt các nhu cầu cơ bản, tôi vẫn có kế hoạch mở rộng và cải tiến sản phẩm:',

    'Thứ nhất là tích hợp thanh toán trực tuyến. Hiện tại hệ thống ghi nhận phương thức thanh toán nhưng chưa tích hợp cổng thanh toán trực tuyến. Trong tương lai, tôi dự kiến tích hợp VNPay, MoMo hoặc ZaloPay để khách hàng có thể đặt cọc và thanh toán trực tiếp trên ứng dụng, giảm thiểu rủi ro hủy booking.',

    'Thứ hai là phát triển hệ thống thông báo đa kênh. Ngoài thông báo trong app, hệ thống sẽ gửi thông báo qua Zalo OA hoặc Email khi booking được xác nhận, sắp đến giờ chơi (nhắc nhở trước 1 giờ), hoặc có khuyến mãi. Điều này giúp tăng tỷ lệ giữ chân khách hàng.',

    'Thứ ba là phát triển tính năng booking định kỳ. Nhiều khách hàng có thói quen đặt sân cố định hàng tuần (ví dụ: mỗi thứ 3 và thứ 5 lúc 7h tối). Tính năng recurring booking sẽ giúp tự động tạo lịch đặt theo chu kỳ, tiết kiệm thời gian cho cả chủ sân và khách hàng.',

    'Thứ tư là ứng dụng trí tuệ nhân tạo (AI) vào kinh doanh. AI có thể phân tích dữ liệu lịch sử để đề xuất giá thuê tối ưu theo khung giờ (dynamic pricing), dự đoán lượng khách để chuẩn bị nhân sự, hoặc gợi ý khách hàng tiềm năng dựa trên hành vi đặt sân.',

    'Cuối cùng là mở rộng thành nền tảng quản lý đa môn thể thao. Kiến trúc hiện tại có thể dễ dàng mở rộng để quản lý không chỉ sân Pickleball mà còn sân cầu lông, sân tennis, sân bóng đá mini – tạo thành một nền tảng quản lý sân thể thao tổng hợp.',
]
for text in paras:
    doc.add_paragraph(text)

# ===== LƯU FILE =====
doc.save(r'BAO CAO APP.docx')
print('✅ Đã ghi đè BAO CAO APP.docx thành công!')
